#!/usr/bin/env python3
"""
YouTube 评论批量抓取 → Word 文档模板
用法: 修改 videos 列表和 output_path，运行即可生成 .docx
然后通过 import_file.sh 上传到腾讯文档
"""
import urllib.request
import json
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

API_KEY = "YOUR_YOUTUBE_API_KEY"

def get_video_comments(video_id, max_results=30):
    """获取视频评论（按相关性排序）"""
    url = f"https://www.googleapis.com/youtube/v3/commentThreads?part=snippet&videoId={video_id}&maxResults={max_results}&key={API_KEY}&order=relevance"
    try:
        with urllib.request.urlopen(url) as response:
            data = json.loads(response.read().decode())
        comments = []
        for item in data.get("items", []):
            snippet = item["snippet"]["topLevelComment"]["snippet"]
            comments.append({
                "author": snippet.get("authorDisplayName", ""),
                "text": snippet.get("textDisplay", ""),
                "likes": snippet.get("likeCount", 0),
                "published": snippet.get("publishedAt", "")
            })
        return comments
    except Exception as e:
        return f"Error: {str(e)}"

def get_video_title(video_id):
    """获取视频标题"""
    url = f"https://www.googleapis.com/youtube/v3/videos?part=snippet&id={video_id}&key={API_KEY}"
    try:
        with urllib.request.urlopen(url) as response:
            data = json.loads(response.read().decode())
        items = data.get("items", [])
        return items[0]["snippet"]["title"] if items else "Unknown Title"
    except Exception as e:
        return f"Error: {str(e)}"

def generate_doc(videos, output_path, title="YouTube 案例视频评论汇总", max_comments=30):
    """生成 Word 文档"""
    doc = Document()
    doc_title = doc.add_heading(title, 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    intro = doc.add_paragraph()
    intro.add_run('文档说明：').bold = True
    intro.add_run(f'共 {len(videos)} 个视频，每个视频最多 {max_comments} 条评论')
    doc.add_paragraph('')

    success_count = 0
    total_comments = 0

    for i, video_id in enumerate(videos):
        print(f"[{i+1}/{len(videos)}] 处理视频: {video_id}")
        video_title = get_video_title(video_id)
        comments = get_video_comments(video_id, max_results=max_comments)

        doc.add_heading(f'{i+1}. {video_title}', level=2)
        p = doc.add_paragraph()
        p.add_run(f'视频链接：https://www.youtube.com/watch?v={video_id}').italic = True

        if isinstance(comments, list) and len(comments) > 0:
            success_count += 1
            total_comments += len(comments)
            doc.add_paragraph(f'评论数量：{len(comments)} 条')
            doc.add_paragraph('')
            for j, comment in enumerate(comments):
                p = doc.add_paragraph()
                p.add_run(f'评论 {j+1}：').bold = True
                p.add_run(f'{comment["author"]} (👍 {comment["likes"]})')
                comment_p = doc.add_paragraph()
                comment_p.add_run(comment["text"])
                doc.add_paragraph('')
        else:
            doc.add_paragraph(f'无法获取评论：{comments}')

        doc.add_paragraph('')
        time.sleep(0.1)

    doc.save(output_path)
    print(f"\n✅ 完成！{success_count}/{len(videos)} 个视频获取到评论")
    print(f"📊 总评论数: {total_comments}")
    print(f"📄 文件: {output_path}")
    return success_count, total_comments

# ============================================================
# 使用示例：修改下方 videos 列表和 output_path
# ============================================================
if __name__ == "__main__":
    videos = [
        "VIDEO_ID_1",
        "VIDEO_ID_2",
        # ... 添加更多视频 ID
    ]
    output_path = "/tmp/YouTube_案例视频评论.docx"
    generate_doc(videos, output_path, max_comments=30)

    # 上传到腾讯文档（取消注释使用）
    # cd ~/.hermes/skills/tencent-docs && bash import_file.sh /tmp/YouTube_案例视频评论.docx
    # 然后 manage.async_import + manage.move_file 移动到目标文件夹
