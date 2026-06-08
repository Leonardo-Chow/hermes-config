"""KOL优秀视频解析 Word文档生成模板"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def create_kol_analysis_doc(
    video_url: str,
    video_title: str,
    blogger_name: str,
    video_summary: str,
    upload_date: str,
    view_count: str,
    duration: str,
    highlights: list[dict],  # [{"title": str, "content": str}]
    blogger_strengths: list[dict],  # [{"title": str, "content": str}]
    screenshot_suggestions: list[dict],  # [{"time": str, "content": str, "purpose": str}]
    summary: str,
    output_path: str
) -> str:
    """生成KOL优秀视频解析Word文档
    
    Args:
        video_url: 视频链接
        video_title: 视频标题
        blogger_name: 博主名称
        video_summary: 视频概述（2-3句话）
        upload_date: 上线时间（YYYY年M月D日格式）
        view_count: 观看次数
        duration: 视频时长
        highlights: 核心亮点列表
        blogger_strengths: 博主优秀之处列表
        screenshot_suggestions: 截图建议列表
        summary: 总结
        output_path: 输出文件路径
        
    Returns:
        输出文件路径
    """
    doc = Document()
    
    # 标题
    title = doc.add_heading('KOL优秀视频解析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 视频基本信息
    doc.add_heading('视频基本信息', level=1)
    
    info_items = [
        ('✅视频链接：', video_url),
        ('✅视频标题：', video_title),
        ('✅博主：', blogger_name),
        ('✅视频概述：', video_summary),
        ('✅上线时间：', upload_date),
        ('✅观看次数：', view_count),
        ('✅视频时长：', duration),
    ]
    
    for label, value in info_items:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(value)
    
    # 核心亮点分析
    doc.add_heading('核心亮点分析', level=1)
    for i, highlight in enumerate(highlights, 1):
        doc.add_heading(f"{i}. {highlight['title']}", level=2)
        p = doc.add_paragraph()
        p.add_run(highlight['content'])
    
    # 博主优秀之处
    doc.add_heading('博主优秀之处', level=1)
    for i, strength in enumerate(blogger_strengths, 1):
        doc.add_heading(f"{i}. {strength['title']}", level=2)
        p = doc.add_paragraph()
        p.add_run(strength['content'])
    
    # 截图建议
    doc.add_heading('截图建议', level=1)
    table = doc.add_table(rows=len(screenshot_suggestions) + 1, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '时间点'
    header_cells[1].text = '截图内容'
    header_cells[2].text = '截图目的'
    
    # 数据行
    for i, suggestion in enumerate(screenshot_suggestions):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = suggestion['time']
        row_cells[1].text = suggestion['content']
        row_cells[2].text = suggestion['purpose']
    
    # 设置表格居中
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 总结
    doc.add_heading('总结', level=1)
    p = doc.add_paragraph()
    p.add_run(summary)
    
    # 保存
    doc.save(output_path)
    return output_path


# 使用示例
if __name__ == '__main__':
    create_kol_analysis_doc(
        video_url='https://www.youtube.com/watch?v=VIDEO_ID',
        video_title='Video Title',
        blogger_name='Blogger Name',
        video_summary='视频概述...',
        upload_date='2026年6月5日',
        view_count='1000次',
        duration='3分30秒',
        highlights=[
            {'title': '功能亮点1', 'content': '详细描述...'},
        ],
        blogger_strengths=[
            {'title': '博主优势1', 'content': '详细描述...'},
        ],
        screenshot_suggestions=[
            {'time': '00:20', 'content': '截图内容', 'purpose': '截图目的'},
        ],
        summary='总结内容...',
        output_path='/tmp/KOL优秀视频解析.docx'
    )
